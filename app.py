import streamlit as st
import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
import os
from fpdf import FPDF
from docx import Document
from io import BytesIO
from openai import OpenAI

# ------------------- OpenAI Summarization -------------------
client = OpenAI(api_key="Your_api_key")

def openai_summarize(text):
    try:
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a data analyst who writes clear, compelling dataset summaries."},
                {"role": "user", "content": f"Please summarize this dataset analysis in detail:\n\n{text}"}
            ],
            temperature=0.3,
            max_tokens=500
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.warning(f"OpenAI summarization failed: {e}")
        return text

# ------------------- Dataset Summary -------------------
def create_generic_dataset_summary(df):
    num_rows, num_cols = df.shape
    dtypes = df.dtypes.to_dict()

    summary_lines = [
        f"The dataset contains {num_rows} rows and {num_cols} columns.",
        "Columns and data types:"
    ]
    for col, dtype in dtypes.items():
        summary_lines.append(f"- {col}: {dtype}")

    stats = df.describe(include='all').round(2).to_string()
    datetime_cols = [col for col in df.columns if pd.api.types.is_datetime64_any_dtype(df[col])]
    time_range_info = ""
    if datetime_cols:
        time_ranges = []
        for col in datetime_cols:
            min_date = df[col].min().strftime("%Y-%m-%d")
            max_date = df[col].max().strftime("%Y-%m-%d")
            time_ranges.append(f"- {col} spans from {min_date} to {max_date}")
        time_range_info = "\n\nDate Ranges:\n" + "\n".join(time_ranges)

    numeric_cols = df.select_dtypes(include='number').columns
    top_values_info = ""
    if len(numeric_cols) > 0:
        top_values_info = "\n\nTop 5 highest summed values per numeric column:\n"
        for col in numeric_cols:
            if not df[col].isna().all():
                top_vals = df[col].sort_values(ascending=False).head(5).round(2)
                top_vals_text = ", ".join(str(v) for v in top_vals)
                top_values_info += f"- {col}: {top_vals_text}\n"

    negative_values_info = ""
    if len(numeric_cols) > 0:
        neg_counts = []
        for col in numeric_cols:
            neg_count = (df[col] < 0).sum()
            if neg_count > 0:
                neg_counts.append(f"- {col}: {neg_count} negative values")
        if neg_counts:
            negative_values_info = "\n\nColumns with negative values:\n" + "\n".join(neg_counts)

    categorical_cols = df.select_dtypes(include='object').columns
    unique_values_info = ""
    if len(categorical_cols) > 0:
        unique_values_info = "\n\nUnique value counts for categorical columns:\n"
        for col in categorical_cols:
            unique_count = df[col].nunique()
            unique_values_info += f"- {col}: {unique_count} unique values\n"

    summary_text = (
        "\n".join(summary_lines)
        + "\n\nSummary Statistics:\n" + stats
        + time_range_info
        + top_values_info
        + negative_values_info
        + unique_values_info
    )
    return summary_text

# ------------------- Plot Generation -------------------
def plot_basic_eda_and_save(df, output_dir="plots"):
    os.makedirs(output_dir, exist_ok=True)
    plot_paths = []

    numeric_cols = df.select_dtypes(include='number')
    categorical_cols = df.select_dtypes(include='object').columns
    datetime_cols = [col for col in df.columns if pd.api.types.is_datetime64_any_dtype(df[col])]

    # Correlation Heatmap
    if numeric_cols.shape[1] > 1:
        plt.figure(figsize=(10,8))
        sns.heatmap(numeric_cols.corr(), annot=True, cmap="coolwarm", fmt=".2f")
        plt.title("Correlation Heatmap")
        path = os.path.join(output_dir, "correlation_heatmap.png")
        plt.savefig(path)
        plt.close()
        plot_paths.append(path)

    # Missing Values Barplot
    missing = df.isnull().sum()
    missing = missing[missing > 0]
    if not missing.empty:
        plt.figure(figsize=(10, 5))
        sns.barplot(x=missing.index, y=missing.values)
        plt.title("Missing Values Count per Column")
        plt.ylabel("Missing Count")
        plt.xticks(rotation=45)
        path = os.path.join(output_dir, "missing_values_barplot.png")
        plt.savefig(path)
        plt.close()
        plot_paths.append(path)

    # Histograms with safe KDE
    for col in numeric_cols.columns:
        plt.figure(figsize=(8,4))
        if len(df) > 50000:
            sns.histplot(df[col].dropna(), kde=False, bins=50)
        else:
            sns.histplot(df[col].dropna(), kde=True)
        plt.title(f"Distribution of {col}")
        path = os.path.join(output_dir, f"hist_{col}.png")
        plt.savefig(path)
        plt.close()
        plot_paths.append(path)

    # Countplots
    for col in categorical_cols:
        plt.figure(figsize=(10,4))
        sns.countplot(y=col, data=df, order=df[col].value_counts().index)
        plt.title(f"Counts of {col}")
        path = os.path.join(output_dir, f"countplot_{col}.png")
        plt.savefig(path)
        plt.close()
        plot_paths.append(path)

    # Time Series Plot
    if datetime_cols and numeric_cols.shape[1] > 0:
        date_col = datetime_cols[0]
        num_col = numeric_cols.columns[0]
        df_time = df[[date_col, num_col]].dropna()
        df_time = df_time.groupby(date_col)[num_col].sum().reset_index()
        plt.figure(figsize=(12,6))
        sns.lineplot(x=date_col, y=num_col, data=df_time)
        plt.title(f"{num_col} over Time")
        path = os.path.join(output_dir, f"time_series_{num_col}.png")
        plt.savefig(path)
        plt.close()
        plot_paths.append(path)

    return plot_paths

# ------------------- PDF -------------------
def create_pdf_report(summary_text, plot_files):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, "Automated Data Analysis Report", ln=True, align="C")
    pdf.ln(10)

    pdf.set_font("Arial", size=12)
    for line in summary_text.split("\n"):
        pdf.multi_cell(0, 8, line)
    pdf.ln(10)

    for plot_path in plot_files:
        pdf.add_page()
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 10, os.path.basename(plot_path), ln=True)
        pdf.image(plot_path, x=10, y=25, w=pdf.w - 20)

    # Return bytes as BytesIO
    pdf_bytes = pdf.output(dest="S").encode("latin1")
    buffer = BytesIO(pdf_bytes)
    return buffer

# ------------------- Word Doc -------------------
def create_word_report(summary_text):
    doc = Document()
    doc.add_heading("Automated Data Analysis Report", level=1)
    doc.add_paragraph(summary_text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ------------------- Streamlit App -------------------
def main():
    st.set_page_config(page_title="CSV to Report Generator", layout="wide")
    st.title("📊 CSV to Report Generator with GPT-4 Summarization")

    uploaded_file = st.file_uploader("Upload your CSV file", type=["csv"])
    if uploaded_file is not None:
        df = pd.read_csv(uploaded_file)

        st.subheader("🔍 Data Preview")
        st.dataframe(df.head())

        raw_summary = create_generic_dataset_summary(df)
        with st.spinner("Generating AI Summary..."):
            summary_text = openai_summarize(raw_summary)

        st.subheader("📝 Detailed Summary")
        st.text_area("Summary Text", summary_text, height=300)

        output_format = st.radio("Select Report Format:", ("PDF", "Word Document"))

        if st.button("Generate Report"):
            plot_files = plot_basic_eda_and_save(df)
            if output_format == "PDF":
                file_buffer = create_pdf_report(summary_text, plot_files)
                st.download_button(
                    label="📄 Download PDF Report",
                    data=file_buffer,
                    file_name="EDA_Report.pdf",
                    mime="application/pdf"
                )
            else:
                file_buffer = create_word_report(summary_text)
                st.download_button(
                    label="📝 Download Word Report",
                    data=file_buffer,
                    file_name="EDA_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

if __name__ == "__main__":
    main()
